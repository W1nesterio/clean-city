from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
INPUT = ROOT / "report_input.docx"
OUTPUT = ROOT / "report_output.docx"


def set_paragraph_text(paragraph, text):
    """Replace paragraph text while keeping the paragraph/run formatting close to the original."""
    if paragraph.runs:
        first = paragraph.runs[0]
        first.text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_font(paragraph, name="Times New Roman", size=Pt(14), bold=None):
    for run in paragraph.runs:
        run.font.name = name
        run.font.size = size
        if bold is not None:
            run.bold = bold
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:ascii"), name)
        r_fonts.set(qn("w:hAnsi"), name)
        r_fonts.set(qn("w:cs"), name)


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def add_code_line_after(paragraph, text):
    p = insert_paragraph_after(paragraph, text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    set_font(p, "Courier New", Pt(8.5), False)
    return p


def merge_adjacent_reference_blocks(text):
    pattern = re.compile(r"\[([0-9,\s]+)\]\.\s+\[([0-9,\s]+)\]")

    def repl(match):
        numbers = []
        for part in (match.group(1), match.group(2)):
            for value in re.findall(r"\d+", part):
                number = int(value)
                if number not in numbers:
                    numbers.append(number)
        numbers.sort()
        return "[" + ", ".join(str(number) for number in numbers) + "]"

    return pattern.sub(repl, text)


SQL_CODE = r"""-- SQL-код создания основных таблиц базы данных приложения "Чистый город"

CREATE TABLE cities (
    id BIGINT UNSIGNED AUTO_INCREMENT PRIMARY KEY,
    name VARCHAR(255) NOT NULL,
    region VARCHAR(255) NULL,
    active TINYINT(1) NOT NULL DEFAULT 1,
    created_at TIMESTAMP NULL,
    updated_at TIMESTAMP NULL
);

CREATE TABLE organizations (
    id BIGINT UNSIGNED AUTO_INCREMENT PRIMARY KEY,
    name VARCHAR(255) NOT NULL,
    city_id BIGINT UNSIGNED NULL,
    district VARCHAR(255) NULL,
    address VARCHAR(255) NULL,
    lat DECIMAL(10,7) NULL,
    lng DECIMAL(10,7) NULL,
    contact_info VARCHAR(255) NULL,
    active TINYINT(1) NOT NULL DEFAULT 1,
    created_at TIMESTAMP NULL,
    updated_at TIMESTAMP NULL,
    INDEX organizations_city_id_index (city_id),
    CONSTRAINT organizations_city_id_fk
        FOREIGN KEY (city_id) REFERENCES cities(id) ON DELETE SET NULL
);

CREATE TABLE users (
    id BIGINT UNSIGNED AUTO_INCREMENT PRIMARY KEY,
    name VARCHAR(255) NOT NULL,
    email VARCHAR(255) NOT NULL UNIQUE,
    email_verified_at TIMESTAMP NULL,
    password VARCHAR(255) NOT NULL,
    role VARCHAR(255) NOT NULL DEFAULT 'resident',
    organization_id BIGINT UNSIGNED NULL,
    avatar_path VARCHAR(255) NULL,
    banned_at TIMESTAMP NULL,
    ban_reason VARCHAR(255) NULL,
    banned_by_user_id BIGINT UNSIGNED NULL,
    points_balance INT NOT NULL DEFAULT 0,
    remember_token VARCHAR(100) NULL,
    created_at TIMESTAMP NULL,
    updated_at TIMESTAMP NULL,
    INDEX users_organization_id_index (organization_id),
    INDEX users_banned_by_user_id_index (banned_by_user_id),
    CONSTRAINT users_organization_id_fk
        FOREIGN KEY (organization_id) REFERENCES organizations(id) ON DELETE SET NULL,
    CONSTRAINT users_banned_by_user_id_fk
        FOREIGN KEY (banned_by_user_id) REFERENCES users(id) ON DELETE SET NULL
);

CREATE TABLE categories (
    id BIGINT UNSIGNED AUTO_INCREMENT PRIMARY KEY,
    name VARCHAR(255) NOT NULL,
    icon VARCHAR(255) NULL,
    active TINYINT(1) NOT NULL DEFAULT 1,
    created_at TIMESTAMP NULL,
    updated_at TIMESTAMP NULL
);

CREATE TABLE tickets (
    id BIGINT UNSIGNED AUTO_INCREMENT PRIMARY KEY,
    user_id BIGINT UNSIGNED NOT NULL,
    category_id BIGINT UNSIGNED NOT NULL,
    assigned_org_id BIGINT UNSIGNED NULL,
    assigned_worker_id BIGINT UNSIGNED NULL,
    status VARCHAR(255) NOT NULL DEFAULT 'created',
    priority VARCHAR(255) NOT NULL DEFAULT 'normal',
    available_to_residents TINYINT(1) NOT NULL DEFAULT 0,
    lat DECIMAL(10,7) NOT NULL,
    lng DECIMAL(10,7) NOT NULL,
    address_text VARCHAR(255) NULL,
    description VARCHAR(200) NULL,
    closed_at TIMESTAMP NULL,
    deleted_by_user_id BIGINT UNSIGNED NULL,
    deleted_at TIMESTAMP NULL,
    delete_reason VARCHAR(255) NULL,
    created_at TIMESTAMP NULL,
    updated_at TIMESTAMP NULL,
    INDEX tickets_user_id_index (user_id),
    INDEX tickets_category_id_index (category_id),
    INDEX tickets_assigned_org_id_index (assigned_org_id),
    INDEX tickets_assigned_worker_id_index (assigned_worker_id),
    INDEX tickets_status_index (status),
    INDEX tickets_priority_index (priority),
    INDEX tickets_deleted_by_user_id_index (deleted_by_user_id),
    INDEX tickets_deleted_at_index (deleted_at),
    CONSTRAINT tickets_user_id_fk
        FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE,
    CONSTRAINT tickets_category_id_fk
        FOREIGN KEY (category_id) REFERENCES categories(id),
    CONSTRAINT tickets_assigned_org_id_fk
        FOREIGN KEY (assigned_org_id) REFERENCES organizations(id) ON DELETE SET NULL,
    CONSTRAINT tickets_assigned_worker_id_fk
        FOREIGN KEY (assigned_worker_id) REFERENCES users(id) ON DELETE SET NULL,
    CONSTRAINT tickets_deleted_by_user_id_fk
        FOREIGN KEY (deleted_by_user_id) REFERENCES users(id) ON DELETE SET NULL
);

CREATE TABLE ticket_photos (
    id BIGINT UNSIGNED AUTO_INCREMENT PRIMARY KEY,
    ticket_id BIGINT UNSIGNED NOT NULL,
    type VARCHAR(255) NOT NULL,
    path VARCHAR(255) NOT NULL,
    created_at TIMESTAMP NULL,
    updated_at TIMESTAMP NULL,
    INDEX ticket_photos_ticket_id_index (ticket_id),
    INDEX ticket_photos_type_index (type),
    CONSTRAINT ticket_photos_ticket_id_fk
        FOREIGN KEY (ticket_id) REFERENCES tickets(id) ON DELETE CASCADE
);

CREATE TABLE ticket_status_histories (
    id BIGINT UNSIGNED AUTO_INCREMENT PRIMARY KEY,
    ticket_id BIGINT UNSIGNED NOT NULL,
    old_status VARCHAR(255) NULL,
    new_status VARCHAR(255) NOT NULL,
    changed_by_user_id BIGINT UNSIGNED NULL,
    comment VARCHAR(255) NULL,
    created_at TIMESTAMP NULL,
    updated_at TIMESTAMP NULL,
    INDEX ticket_status_histories_ticket_id_index (ticket_id),
    INDEX ticket_status_histories_changed_by_user_id_index (changed_by_user_id),
    INDEX ticket_status_histories_new_status_index (new_status),
    CONSTRAINT ticket_status_histories_ticket_id_fk
        FOREIGN KEY (ticket_id) REFERENCES tickets(id) ON DELETE CASCADE,
    CONSTRAINT ticket_status_histories_changed_by_user_id_fk
        FOREIGN KEY (changed_by_user_id) REFERENCES users(id) ON DELETE SET NULL
);

CREATE TABLE api_tokens (
    id BIGINT UNSIGNED AUTO_INCREMENT PRIMARY KEY,
    user_id BIGINT UNSIGNED NOT NULL,
    name VARCHAR(255) NOT NULL DEFAULT 'api-token',
    token_hash VARCHAR(64) NOT NULL UNIQUE,
    last_used_at TIMESTAMP NULL,
    expires_at TIMESTAMP NULL,
    created_at TIMESTAMP NULL,
    updated_at TIMESTAMP NULL,
    INDEX api_tokens_user_id_index (user_id),
    CONSTRAINT api_tokens_user_id_fk
        FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
);

CREATE TABLE worker_registration_codes (
    id BIGINT UNSIGNED AUTO_INCREMENT PRIMARY KEY,
    code VARCHAR(32) NOT NULL UNIQUE,
    issued_to VARCHAR(120) NULL,
    organization_id BIGINT UNSIGNED NOT NULL,
    created_by_user_id BIGINT UNSIGNED NULL,
    used_by_user_id BIGINT UNSIGNED NULL,
    used_at TIMESTAMP NULL,
    max_uses INT UNSIGNED NOT NULL DEFAULT 1,
    used_count INT UNSIGNED NOT NULL DEFAULT 0,
    expires_at TIMESTAMP NULL,
    active TINYINT(1) NOT NULL DEFAULT 1,
    revoked_at TIMESTAMP NULL,
    revoked_by_user_id BIGINT UNSIGNED NULL,
    created_at TIMESTAMP NULL,
    updated_at TIMESTAMP NULL,
    INDEX worker_registration_codes_organization_id_index (organization_id),
    INDEX worker_registration_codes_created_by_user_id_index (created_by_user_id),
    INDEX worker_registration_codes_used_by_user_id_index (used_by_user_id),
    INDEX worker_registration_codes_revoked_by_user_id_index (revoked_by_user_id),
    INDEX worker_registration_codes_active_index (active),
    INDEX worker_registration_codes_expires_at_index (expires_at),
    CONSTRAINT worker_registration_codes_organization_id_fk
        FOREIGN KEY (organization_id) REFERENCES organizations(id) ON DELETE CASCADE,
    CONSTRAINT worker_registration_codes_created_by_user_id_fk
        FOREIGN KEY (created_by_user_id) REFERENCES users(id) ON DELETE SET NULL,
    CONSTRAINT worker_registration_codes_used_by_user_id_fk
        FOREIGN KEY (used_by_user_id) REFERENCES users(id) ON DELETE SET NULL,
    CONSTRAINT worker_registration_codes_revoked_by_user_id_fk
        FOREIGN KEY (revoked_by_user_id) REFERENCES users(id) ON DELETE SET NULL
);

CREATE TABLE organization_user_blocks (
    id BIGINT UNSIGNED AUTO_INCREMENT PRIMARY KEY,
    organization_id BIGINT UNSIGNED NOT NULL,
    user_id BIGINT UNSIGNED NOT NULL,
    created_by_user_id BIGINT UNSIGNED NULL,
    reason VARCHAR(255) NULL,
    active TINYINT(1) NOT NULL DEFAULT 1,
    created_at TIMESTAMP NULL,
    updated_at TIMESTAMP NULL,
    INDEX organization_user_blocks_organization_id_index (organization_id),
    INDEX organization_user_blocks_user_id_index (user_id),
    INDEX organization_user_blocks_created_by_user_id_index (created_by_user_id),
    INDEX organization_user_blocks_active_index (active),
    CONSTRAINT organization_user_blocks_organization_id_fk
        FOREIGN KEY (organization_id) REFERENCES organizations(id) ON DELETE CASCADE,
    CONSTRAINT organization_user_blocks_user_id_fk
        FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE,
    CONSTRAINT organization_user_blocks_created_by_user_id_fk
        FOREIGN KEY (created_by_user_id) REFERENCES users(id) ON DELETE SET NULL
);

CREATE TABLE ticket_hides (
    id BIGINT UNSIGNED AUTO_INCREMENT PRIMARY KEY,
    ticket_id BIGINT UNSIGNED NOT NULL,
    organization_id BIGINT UNSIGNED NULL,
    hidden_by_user_id BIGINT UNSIGNED NULL,
    reason VARCHAR(255) NULL,
    active TINYINT(1) NOT NULL DEFAULT 1,
    created_at TIMESTAMP NULL,
    updated_at TIMESTAMP NULL,
    INDEX ticket_hides_ticket_id_index (ticket_id),
    INDEX ticket_hides_organization_id_index (organization_id),
    INDEX ticket_hides_hidden_by_user_id_index (hidden_by_user_id),
    INDEX ticket_hides_active_index (active),
    CONSTRAINT ticket_hides_ticket_id_fk
        FOREIGN KEY (ticket_id) REFERENCES tickets(id) ON DELETE CASCADE,
    CONSTRAINT ticket_hides_organization_id_fk
        FOREIGN KEY (organization_id) REFERENCES organizations(id) ON DELETE SET NULL,
    CONSTRAINT ticket_hides_hidden_by_user_id_fk
        FOREIGN KEY (hidden_by_user_id) REFERENCES users(id) ON DELETE SET NULL
);

CREATE TABLE organization_complaints (
    id BIGINT UNSIGNED AUTO_INCREMENT PRIMARY KEY,
    organization_id BIGINT UNSIGNED NOT NULL,
    created_by_user_id BIGINT UNSIGNED NOT NULL,
    target_user_id BIGINT UNSIGNED NULL,
    ticket_id BIGINT UNSIGNED NULL,
    type VARCHAR(40) NOT NULL DEFAULT 'other',
    status VARCHAR(30) NOT NULL DEFAULT 'pending',
    title VARCHAR(160) NOT NULL,
    description TEXT NULL,
    reviewed_by_user_id BIGINT UNSIGNED NULL,
    reviewed_at TIMESTAMP NULL,
    resolution TEXT NULL,
    created_at TIMESTAMP NULL,
    updated_at TIMESTAMP NULL,
    INDEX organization_complaints_organization_id_index (organization_id),
    INDEX organization_complaints_created_by_user_id_index (created_by_user_id),
    INDEX organization_complaints_target_user_id_index (target_user_id),
    INDEX organization_complaints_ticket_id_index (ticket_id),
    INDEX organization_complaints_type_index (type),
    INDEX organization_complaints_status_index (status),
    INDEX organization_complaints_reviewed_by_user_id_index (reviewed_by_user_id),
    CONSTRAINT organization_complaints_organization_id_fk
        FOREIGN KEY (organization_id) REFERENCES organizations(id) ON DELETE CASCADE,
    CONSTRAINT organization_complaints_created_by_user_id_fk
        FOREIGN KEY (created_by_user_id) REFERENCES users(id) ON DELETE CASCADE,
    CONSTRAINT organization_complaints_target_user_id_fk
        FOREIGN KEY (target_user_id) REFERENCES users(id) ON DELETE SET NULL,
    CONSTRAINT organization_complaints_ticket_id_fk
        FOREIGN KEY (ticket_id) REFERENCES tickets(id) ON DELETE SET NULL,
    CONSTRAINT organization_complaints_reviewed_by_user_id_fk
        FOREIGN KEY (reviewed_by_user_id) REFERENCES users(id) ON DELETE SET NULL
);

CREATE TABLE ticket_claim_requests (
    id BIGINT UNSIGNED AUTO_INCREMENT PRIMARY KEY,
    ticket_id BIGINT UNSIGNED NOT NULL,
    worker_id BIGINT UNSIGNED NOT NULL,
    organization_id BIGINT UNSIGNED NOT NULL,
    status VARCHAR(30) NOT NULL DEFAULT 'pending',
    comment VARCHAR(255) NULL,
    reviewed_by_user_id BIGINT UNSIGNED NULL,
    reviewed_at TIMESTAMP NULL,
    resolution TEXT NULL,
    created_at TIMESTAMP NULL,
    updated_at TIMESTAMP NULL,
    INDEX ticket_claim_requests_ticket_id_index (ticket_id),
    INDEX ticket_claim_requests_worker_id_index (worker_id),
    INDEX ticket_claim_requests_organization_id_index (organization_id),
    INDEX ticket_claim_requests_status_index (status),
    CONSTRAINT ticket_claim_requests_ticket_id_fk
        FOREIGN KEY (ticket_id) REFERENCES tickets(id) ON DELETE CASCADE,
    CONSTRAINT ticket_claim_requests_worker_id_fk
        FOREIGN KEY (worker_id) REFERENCES users(id) ON DELETE CASCADE,
    CONSTRAINT ticket_claim_requests_organization_id_fk
        FOREIGN KEY (organization_id) REFERENCES organizations(id) ON DELETE CASCADE,
    CONSTRAINT ticket_claim_requests_reviewed_by_user_id_fk
        FOREIGN KEY (reviewed_by_user_id) REFERENCES users(id) ON DELETE SET NULL
);

CREATE TABLE news (
    id BIGINT UNSIGNED AUTO_INCREMENT PRIMARY KEY,
    title VARCHAR(255) NOT NULL,
    body TEXT NULL,
    published_date DATE NOT NULL,
    active TINYINT(1) NOT NULL DEFAULT 1,
    created_by_user_id BIGINT UNSIGNED NULL,
    organization_id BIGINT UNSIGNED NULL,
    created_at TIMESTAMP NULL,
    updated_at TIMESTAMP NULL,
    INDEX news_created_by_user_id_index (created_by_user_id),
    INDEX news_organization_id_index (organization_id),
    CONSTRAINT news_created_by_user_id_fk
        FOREIGN KEY (created_by_user_id) REFERENCES users(id) ON DELETE SET NULL,
    CONSTRAINT news_organization_id_fk
        FOREIGN KEY (organization_id) REFERENCES organizations(id) ON DELETE SET NULL
);

CREATE TABLE news_photos (
    id BIGINT UNSIGNED AUTO_INCREMENT PRIMARY KEY,
    news_id BIGINT UNSIGNED NOT NULL,
    path VARCHAR(255) NOT NULL,
    sort_order INT NOT NULL DEFAULT 0,
    created_at TIMESTAMP NULL,
    updated_at TIMESTAMP NULL,
    CONSTRAINT news_photos_news_id_fk
        FOREIGN KEY (news_id) REFERENCES news(id) ON DELETE CASCADE
);

CREATE TABLE rewards (
    id BIGINT UNSIGNED AUTO_INCREMENT PRIMARY KEY,
    title VARCHAR(255) NOT NULL,
    description TEXT NULL,
    photo_path VARCHAR(255) NULL,
    points_required INT NOT NULL DEFAULT 0,
    code VARCHAR(255) NULL,
    valid_from DATE NULL,
    valid_to DATE NULL,
    active TINYINT(1) NOT NULL DEFAULT 1,
    created_by_user_id BIGINT UNSIGNED NULL,
    organization_id BIGINT UNSIGNED NULL,
    created_at TIMESTAMP NULL,
    updated_at TIMESTAMP NULL,
    INDEX rewards_created_by_user_id_index (created_by_user_id),
    INDEX rewards_organization_id_index (organization_id),
    CONSTRAINT rewards_created_by_user_id_fk
        FOREIGN KEY (created_by_user_id) REFERENCES users(id) ON DELETE SET NULL,
    CONSTRAINT rewards_organization_id_fk
        FOREIGN KEY (organization_id) REFERENCES organizations(id) ON DELETE SET NULL
);

CREATE TABLE points_transactions (
    id BIGINT UNSIGNED AUTO_INCREMENT PRIMARY KEY,
    user_id BIGINT UNSIGNED NOT NULL,
    amount INT NOT NULL,
    balance_after INT NOT NULL,
    reason VARCHAR(255) NULL,
    admin_id BIGINT UNSIGNED NULL,
    created_at TIMESTAMP NULL,
    updated_at TIMESTAMP NULL,
    INDEX points_transactions_user_id_index (user_id),
    INDEX points_transactions_admin_id_index (admin_id),
    CONSTRAINT points_transactions_user_id_fk
        FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE,
    CONSTRAINT points_transactions_admin_id_fk
        FOREIGN KEY (admin_id) REFERENCES users(id) ON DELETE SET NULL
);

CREATE TABLE system_settings (
    id BIGINT UNSIGNED AUTO_INCREMENT PRIMARY KEY,
    `key` VARCHAR(255) NOT NULL UNIQUE,
    value TEXT NULL,
    description VARCHAR(255) NULL,
    created_at TIMESTAMP NULL,
    updated_at TIMESTAMP NULL
);"""


doc = Document(INPUT)

bad_ref_pattern = re.compile(r"([.!?])\s+(\[[0-9,\s]+\])")
fixed_refs = 0

for paragraph in doc.paragraphs:
    original = paragraph.text
    changed = merge_adjacent_reference_blocks(bad_ref_pattern.sub(r" \2\1", original))
    if changed != original:
        set_paragraph_text(paragraph, changed)
        fixed_refs += 1

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                original = paragraph.text
                changed = merge_adjacent_reference_blocks(bad_ref_pattern.sub(r" \2\1", original))
                if changed != original:
                    set_paragraph_text(paragraph, changed)
                    fixed_refs += 1

appendix_paragraph = None
for paragraph in reversed(doc.paragraphs):
    if paragraph.text.strip().upper() == "ПРИЛОЖЕНИЕ":
        appendix_paragraph = paragraph
        break

if appendix_paragraph is None:
    raise RuntimeError("Не найден заголовок 'ПРИЛОЖЕНИЕ' в конце документа")

set_paragraph_text(appendix_paragraph, "ПРИЛОЖЕНИЕ А")
appendix_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
appendix_paragraph.paragraph_format.first_line_indent = Pt(0)
appendix_paragraph.paragraph_format.space_before = Pt(0)
appendix_paragraph.paragraph_format.space_after = Pt(0)
set_font(appendix_paragraph, "Times New Roman", Pt(14), True)

current = insert_paragraph_after(appendix_paragraph, "Код базы данных")
current.alignment = WD_ALIGN_PARAGRAPH.CENTER
current.paragraph_format.first_line_indent = Pt(0)
current.paragraph_format.space_after = Pt(12)
set_font(current, "Times New Roman", Pt(14), True)

current = insert_paragraph_after(current, "Листинг А.1 - SQL-код создания основных таблиц базы данных")
current.alignment = WD_ALIGN_PARAGRAPH.LEFT
current.paragraph_format.first_line_indent = Pt(0)
current.paragraph_format.space_before = Pt(0)
current.paragraph_format.space_after = Pt(6)
set_font(current, "Times New Roman", Pt(14), False)

for line in SQL_CODE.splitlines():
    current = add_code_line_after(current, line)

doc.save(OUTPUT)
print(f"fixed_refs={fixed_refs}")
print(f"output={OUTPUT}")
